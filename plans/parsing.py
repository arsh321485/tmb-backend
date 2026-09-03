"""
B2: plan parsing -- step one.

Full B2 ("extracts roles, triggers, thresholds, escalation paths, RTO/RPO
targets, notification duties and contact trees into a structured plan
model") needs real NLP/document-understanding work. What's here is the
necessary first step: pulling the raw text out of the DOCX/PDF/XLSX file
so there's something to run that analysis on later. Structured extraction
is intentionally not attempted yet.
"""

import io
import logging

logger = logging.getLogger(__name__)


class PlanParsingError(Exception):
    pass


def extract_text(file_bytes: bytes, extension: str) -> str:
    extension = extension.lower()
    if extension == "docx":
        return _extract_docx(file_bytes)
    if extension == "pdf":
        return _extract_pdf(file_bytes)
    if extension == "xlsx":
        return _extract_xlsx(file_bytes)
    raise PlanParsingError(f"Unsupported extension: {extension}")


def _extract_docx(file_bytes: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for p in document.paragraphs:
        if not p.text.strip():
            continue
        # Mark headings so structured_extraction.py can find sections
        # (e.g. "Escalation", "Contacts") by heading text, not just guesswork.
        if p.style and p.style.name.startswith("Heading"):
            paragraphs.append(f"## {p.text}")
        else:
            paragraphs.append(p.text)

    # Plans often carry key info (RACI grids, contact lists) in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    pages_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return "\n".join(pages_text)


def _extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"# {sheet.title}")
        for row in sheet.iter_rows():
            values = [str(c.value) for c in row if c.value is not None]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)
